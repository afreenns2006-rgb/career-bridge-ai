"""
Resume parsing module for Career Bridge AI.

Handles resume file processing, text extraction, skill identification,
and ATS scoring.
"""

from io import BytesIO
from pathlib import Path
from typing import Optional, List, Dict, Any
import logging
import re
from collections import Counter

try:
    from pypdf import PdfReader as PypdfReader
except ImportError:
    PypdfReader = None

try:
    from PyPDF2 import PdfReader as PyPDF2Reader
except ImportError:
    PyPDF2Reader = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)

RESUME_KEYWORDS = {
    "skills",
    "education",
    "project",
    "experience",
    "python",
    "java",
    "sql",
    "internship",
    "college",
    "degree",
}

# Common skills list for extraction
COMMON_SKILLS = {
    "python",
    "java",
    "javascript",
    "c++",
    "c#",
    "php",
    "ruby",
    "go",
    "rust",
    "kotlin",
    "sql",
    "mongodb",
    "postgresql",
    "mysql",
    "oracle",
    "redis",
    "elasticsearch",
    "react",
    "angular",
    "vue",
    "node.js",
    "django",
    "flask",
    "fastapi",
    "spring",
    "aws",
    "azure",
    "gcp",
    "docker",
    "kubernetes",
    "git",
    "jenkins",
    "terraform",
    "machine learning",
    "deep learning",
    "nlp",
    "computer vision",
    "tensorflow",
    "pytorch",
    "html",
    "css",
    "rest api",
    "graphql",
    "microservices",
    "agile",
    "scrum",
    "communication",
    "leadership",
    "teamwork",
    "problem solving",
    "data analysis",
}


class ResumeParser:
    """
    Parses resume documents and extracts relevant information.

    Supports multiple file formats and provides skill extraction
    and ATS score calculation.
    """

    def __init__(self, file_path: Path | None = None, file_name: str | None = None, file_bytes: bytes | None = None) -> None:
        """
        Initialize resume parser with a file path.

        Args:
            file_path: Path to resume file (PDF, DOCX, TXT).

        TODO: Implement parser initialization.
        """
        self.file_path = file_path
        self.file_name = file_name or (file_path.name if file_path else "uploaded_resume")
        self.file_bytes = file_bytes
        self.text: Optional[str] = None
        self.metadata: dict[str, Any] = {}

    @classmethod
    def from_upload(cls, uploaded_file: Any) -> "ResumeParser":
        """Create a parser from a Streamlit UploadedFile without writing to disk."""
        return cls(
            file_name=getattr(uploaded_file, "name", "uploaded_resume"),
            file_bytes=uploaded_file.getvalue(),
        )

    def extract_text(self) -> str:
        """
        Extract text from resume file.

        Supports PDF, DOCX, and TXT formats. Handles formatting
        and encoding issues.

        Returns:
            Extracted text from resume.

        Raises:
            FileNotFoundError: If resume file not found.
            ValueError: If file format not supported.

        TODO: Implement text extraction from multiple formats.
        """
        if self.file_bytes is None and (self.file_path is None or not self.file_path.exists()):
            raise FileNotFoundError(f"Resume file not found: {self.file_path}")

        file_ext = Path(self.file_name).suffix.lower()

        try:
            if file_ext == ".pdf":
                self.text = self._extract_from_pdf()
            elif file_ext == ".docx":
                self.text = self._extract_from_docx()
            elif file_ext == ".txt":
                self.text = self._extract_from_txt()
            else:
                raise ValueError(f"Unsupported file format: {file_ext}")

            return self.text if self.text else ""
        except Exception as e:
            logger.error(f"Error extracting text from resume: {e}")
            raise

    def _extract_from_pdf(self) -> str:
        """Extract text from PDF file."""
        if not PypdfReader and not PyPDF2Reader and not pdfplumber:
            logger.warning("No PDF reader is installed. Cannot extract PDF.")
            return ""

        attempted_readers = set()
        for reader_name, reader_cls in (("pypdf", PypdfReader), ("PyPDF2", PyPDF2Reader)):
            if reader_cls is None or reader_cls in attempted_readers:
                continue
            attempted_readers.add(reader_cls)
            text = self._extract_pdf_with_reader(reader_cls, reader_name)
            if text.strip():
                if reader_name != "pypdf":
                    self.metadata["used_fallback_parser"] = True
                return text

        text = self._extract_pdf_with_pdfplumber()
        if text.strip():
            self.metadata["used_fallback_parser"] = True
            return text

        return ""

    def _get_binary_stream(self) -> BytesIO:
        """Return upload bytes as a fresh in-memory stream."""
        if self.file_bytes is not None:
            return BytesIO(self.file_bytes)
        with open(self.file_path, "rb") as file:
            return BytesIO(file.read())

    def _extract_pdf_with_reader(self, reader_cls: Any, reader_name: str) -> str:
        """Extract PDF text with pypdf/PyPDF2 from an in-memory stream."""
        text_parts = []
        try:
            with self._get_binary_stream() as file:
                reader = reader_cls(file)
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
        except Exception as e:
            logger.warning("%s PDF extraction failed: %s", reader_name, e)

        return "\n".join(text_parts).strip()

    def _extract_pdf_with_pdfplumber(self) -> str:
        """Extract PDF text with pdfplumber as a deployment-compatible fallback."""
        if not pdfplumber:
            logger.warning("pdfplumber is not installed. Skipping PDF fallback parser.")
            return ""

        text_parts = []
        try:
            with self._get_binary_stream() as file:
                with pdfplumber.open(file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        if page_text.strip():
                            text_parts.append(page_text)
        except Exception as e:
            logger.warning("pdfplumber PDF extraction failed: %s", e)

        return "\n".join(text_parts).strip()

    def _extract_from_docx(self) -> str:
        """Extract text from DOCX file."""
        if not Document:
            logger.warning("python-docx not installed. Cannot extract DOCX.")
            return ""

        text_parts = []
        try:
            doc = Document(BytesIO(self.file_bytes)) if self.file_bytes is not None else Document(self.file_path)
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")

        return "\n".join(text_parts).strip()

    def _extract_from_txt(self) -> str:
        """Extract text from TXT file."""
        try:
            if self.file_bytes is not None:
                try:
                    return self.file_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    return self.file_bytes.decode("latin-1", errors="ignore")
            with open(self.file_path, encoding="utf-8") as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading TXT: {e}")
            return ""

    def extract_skills(self) -> list[str]:
        """
        Extract skills from resume text.

        Uses NLP techniques and predefined skill vocabularies
        to identify technical and soft skills.

        Returns:
            List of identified skills.

        TODO: Implement skill extraction using spaCy/Sentence Transformers.
        """
        if not self.text:
            self.extract_text()

        text_lower = self.text.lower() if self.text else ""
        found_skills = set()

        for skill in COMMON_SKILLS:
            if skill in text_lower:
                found_skills.add(skill)

        return list(found_skills)

    def extract_education(self) -> list[dict[str, str]]:
        """
        Extract education information from resume.

        Returns:
            List of education entries with degree, institution, etc.

        TODO: Implement education extraction.
        """
        if not self.text:
            self.extract_text()

        education_list = []
        education_keywords = ["b.tech", "b.sc", "m.tech", "mca", "mba", "bachelor", "master", "degree"]

        text_lower = self.text.lower() if self.text else ""

        # Simple extraction based on keywords
        for keyword in education_keywords:
            if keyword in text_lower:
                education_list.append({"type": keyword, "mentioned": True})

        return education_list

    def extract_experience(self) -> list[dict[str, str]]:
        """
        Extract work experience from resume.

        Returns:
            List of experience entries with title, company, duration, etc.

        TODO: Implement experience extraction.
        """
        if not self.text:
            self.extract_text()

        experience_list = []

        # Simple extraction based on common job titles
        job_titles = ["developer", "engineer", "analyst", "manager", "consultant", "designer", "lead"]
        text_lower = self.text.lower() if self.text else ""

        for title in job_titles:
            if title in text_lower:
                experience_list.append({"title": title, "mentioned": True})

        return experience_list

    def calculate_ats_score(self) -> float:
        """
        Calculate ATS (Applicant Tracking System) score for resume.

        Evaluates resume formatting, keywords, structure, and content
        to generate a score between 0-100.

        Returns:
            ATS score as float between 0 and 100.

        TODO: Implement ATS scoring algorithm.
        """
        if not self.text:
            self.extract_text()

        score = 0.0
        text = self.text if self.text else ""
        text_lower = text.lower()

        # Content length check
        if len(text) > 500:
            score += 10
        if len(text) > 1000:
            score += 10

        # Skills presence
        skills = self.extract_skills()
        score += min(len(skills) * 2, 20)

        # Education presence
        education = self.extract_education()
        if education:
            score += 15

        # Experience presence
        experience = self.extract_experience()
        if experience:
            score += 15

        # Email and contact info
        email_pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b"
        if re.search(email_pattern, text):
            score += 10

        # Phone number
        phone_pattern = r"\b\d{10}\b|\b\d{3}[-.\s]\d{3}[-.\s]\d{4}\b"
        if re.search(phone_pattern, text):
            score += 5

        # Work keywords
        work_keywords = ["experience", "achieved", "managed", "developed", "implemented"]
        keyword_count = sum(1 for kw in work_keywords if kw in text_lower)
        score += keyword_count * 2

        # Formatting check (multiple lines, structure)
        lines = text.split("\n")
        if len(lines) > 10:
            score += 10

        return min(score, 100.0)

    def get_resume_summary(self) -> dict[str, Any]:
        """
        Get comprehensive resume summary.

        Returns:
            Dictionary containing all extracted information and scores.

        TODO: Implement resume summary compilation.
        """
        if not self.text:
            self.extract_text()

        return {
            "filename": self.file_name,
            "text_length": len(self.text) if self.text else 0,
            "skills": self.extract_skills(),
            "education": self.extract_education(),
            "experience": self.extract_experience(),
            "ats_score": self.calculate_ats_score(),
            "extracted": bool(self.text),
        }

    def validate_resume(self) -> bool:
        """
        Validate resume file and content.

        Checks file format, file size, and content requirements.

        Returns:
            True if resume is valid, False otherwise.

        TODO: Implement validation checks.
        """
        # Check file exists or upload bytes are available.
        if self.file_bytes is None and (self.file_path is None or not self.file_path.exists()):
            return False

        # Check file format
        valid_formats = {".pdf", ".docx", ".txt"}
        if Path(self.file_name).suffix.lower() not in valid_formats:
            return False

        # Check file size (max 10MB)
        max_size = 10 * 1024 * 1024
        file_size = len(self.file_bytes) if self.file_bytes is not None else self.file_path.stat().st_size
        if file_size > max_size:
            return False

        # Try to extract text
        try:
            text = self.extract_text()
            return self.has_enough_resume_text(text)
        except Exception:
            return False

    @staticmethod
    def has_enough_resume_text(text: str) -> bool:
        """Validate extracted text without rejecting deployment PDFs too aggressively."""
        text = text or ""
        text_lower = text.lower()
        return len(text.strip()) >= 30 or any(keyword in text_lower for keyword in RESUME_KEYWORDS)
